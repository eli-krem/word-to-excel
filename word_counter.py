import xlsxwriter
import os
from docx import Document
from os import listdir
from os.path import isfile, join

project_dir = os.path.dirname(os.path.realpath(__file__))
dir = project_dir + "\\data\\"
files = [f for f in listdir(dir) if isfile(join(dir, f))]

data = []

for file in files:
 filename = file
 document = Document(dir + file)
 lines = len(document.paragraphs)
 line_num = 0
 for line in document.paragraphs:
  line_num += 1
  words = line.text.strip()
  words_array = words.split(" ")
  count = len(words_array)
  first_word = words_array[0].replace(",", "").replace(".", "")
  second_word = words_array[1].replace(",", "").replace(".", "") if len(words_array) >= 2 else ""
  last_word = words_array[-1].replace(",", "").replace(".", "")
  file_array = [filename, lines, line_num, count, words, first_word, second_word, last_word]
  data.append(file_array)

workbook = xlsxwriter.Workbook('Results.xlsx')
worksheet = workbook.add_worksheet()
row = 0
col = 0

for filename, lines, line_num, count, words, first_word, second_word, last_word in data:
    worksheet.write(row, col, filename)
    worksheet.write(row, col + 1, lines)
    worksheet.write(row, col + 2, line_num)
    worksheet.write(row, col + 3, count)
    worksheet.write(row, col + 4, words)
    worksheet.write(row, col + 5, first_word)
    worksheet.write(row, col + 6, second_word)
    worksheet.write(row, col + 7, last_word)
    row += 1


workbook.close()

